import hashlib
from pathlib import Path

from services.document_record_service import DocumentRecordService
from services.ocr_service import ocr_service
from services.vector_store_manager import vector_store_manager
from services.document_split_service import document_split_service
from core.config import settings
from db.session import Session
from loguru import logger


class VectorIndexService:
    """
    向量索引服务
    负责读取文件、生成向量、存储到向量数据库
    """

    def index_single_file(self, file_path: str, db: Session):
        """
        索引单个文件

        :param file_path: 文件路径
        :param db: 数据库对象
        """
        try:
            # 1.读取文件内容
            normalized_path, doc_name, doc_hash, content, suffix = self.read_file(file_path)
            # 2.是否存在文件名和hash一致的文件
            same_name_doc = DocumentRecordService.get_by_name(db, doc_name)
            if same_name_doc:
                if same_name_doc.doc_hash == doc_hash:
                    logger.info(f"文件{file_path}已存在，跳过索引")
                    raise RuntimeError("当前知识库已有相同文件存在，无需重复索引")
                else:
                    vector_store_manager.delete_documents_by_name(doc_name)

            # 3.使用分割器
            docs = document_split_service.split_document(content, normalized_path)
            # 4.生成向量并存储
            if docs:
                vector_store_manager.add_documents(docs)
                # 存储文档记录
                DocumentRecordService.save_or_update_document(db, doc_name=doc_name,
                                                              doc_path=file_path, doc_hash=doc_hash,
                                                              doc_type=Path(normalized_path).suffix,
                                                              doc_source="file_upload",
                                                              chunk_count=len(docs), status="active")
            else:
                logger.info(f"文件{file_path}没有可索引的内容")
                raise RuntimeError("当前文件没有可索引的内容")

        except Exception as e:
            logger.error(f"索引文件{file_path}时出错:{e}")
            raise RuntimeError(f"索引文件失败: {e}") from e

    def read_file(self, file_path: str):
        """
        读取文件
        """
        path = Path(file_path).resolve()

        if not path.exists() or not path.is_file():
            raise ValueError(f"文件不存在:{path}")
        logger.info(f"开始索引文件:{path}")
        # 分块读取二进制流 计算hash值
        doc_hash = self.get_hash(path)

        # 根据后缀提取文档内容
        suffix = path.suffix.lower()
        normalized_path = path.as_posix()
        doc_name = Path(normalized_path).name
        content = self.do_read(path, suffix)

        logger.info(f"读取文件{file_path}, 内容长度: {len(content)}字符, 哈希值：{doc_hash}")
        return normalized_path, doc_name, doc_hash, content, suffix

    def get_hash(self, path: Path) -> str:
        hasher = hashlib.md5()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    def do_read(self, path: Path, suffix: str) -> str:
        """
        根据文件类型，按策略读取文件内容
        :param path 文件路径
        :param suffix 文件后缀
        """
        if suffix in [".txt", ".md"]:
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                raise ValueError(f"文件编码非UTF-8")
        elif suffix in [".png", ".jpg", ".jpeg"]:
            content = ocr_service.ocr(path)
        elif suffix == ".pdf":
            content = self.read_from_pdf(path)
        elif suffix == ".docx":
            content = self.read_from_docx(path)
        elif suffix == ".doc":
            content = self.read_from_doc(path)
        elif suffix in [".xlsx", ".xls"]:
            content = self.read_from_excel(path, suffix)
        elif suffix == ".csv":
            content = self.read_from_csv(path)
        else:
            raise ValueError(f"不支持的文件类型:{suffix}")

        return content

    def read_from_pdf(self, path: Path) -> str:
        """
        从pdf文件中读取内容
        """
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text_part = []
        empty_page_count = 0
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_part.append(page_text)
            else:
                # 空页+1
                empty_page_count += 1
        total_pages = len(reader.pages)
        extracted_text = "\n".join(text_part).strip()
        # 判定是否为扫描件：
        # 1. 完全提取不到文本
        # 2. 或超过半数页面无文本，且总文本量低于阈值
        is_scanned = (
                not extracted_text
                or (
                        total_pages > 0
                        and empty_page_count / total_pages >= 0.5
                )
        )
        if not is_scanned:
            if not text_part:
                raise ValueError(f"PDF文件{path}无有效文本")
            return extracted_text

        # 执行OCR扫描
        logger.info(
            f"PDF文件{path}疑似扫描件（共{total_pages}页，"
            f"其中{empty_page_count}页无文本，提取到{len(extracted_text)}字符），"
            f"尝试使用OCR识别"
        )
        return ocr_service.ocr(path)

    def read_from_docx(self, path: Path) -> str:
        """
        从docx文件中读取内容
        """
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text_part = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_part.append(text)
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_text:
                    text_part.append(" ".join(row_text))
        extracted_text = "\n".join(text_part).strip()
        if not extracted_text:
            raise ValueError(f"docx文件{path}无有效文本")
        return extracted_text

    def read_from_doc(self, path: Path) -> str:
        """
        从doc文件中读取内容
        优先用LibreOffice转换为docx后复用read_from_docx（保留表格结构）；
        若LibreOffice未配置或转换失败，降级到antiword纯文本提取。
        """
        # 优先尝试LibreOffice转换
        soffice_path = settings.libreoffice_path
        if soffice_path and Path(soffice_path).exists():
            try:
                converted = self._convert_doc_to_docx(path, soffice_path)
                if converted:
                    content = self.read_from_docx(converted)
                    logger.info(f"doc文件{path}通过LibreOffice转换解析完成")
                    return content
            except Exception as e:
                logger.warning(f"LibreOffice转换失败，降级到antiword: {e}")
        # 降级：antiword纯文本
        return self._read_from_doc_antiword(path)

    def _convert_doc_to_docx(self, path: Path, soffice_path: str) -> Path:
        """
        调用LibreOffice headless将.doc转换为.docx，返回转换后的文件路径
        """
        import subprocess
        import tempfile
        import shutil
        with tempfile.TemporaryDirectory() as tmp_dir:
            result = subprocess.run(
                [soffice_path, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, str(path)],
                capture_output=True,
                timeout=120,
            )
            if result.returncode != 0:
                err = result.stderr.decode("utf-8", errors="ignore").strip()
                raise RuntimeError(f"soffice转换失败: {err}")
            converted = Path(tmp_dir) / (path.stem + ".docx")
            if not converted.exists():
                raise RuntimeError(f"soffice未生成输出文件: {converted}")
            # 复制到稳定位置（临时目录退出后会被清理）；shutil支持跨盘移动
            stable_path = path.with_suffix(".converted.docx")
            shutil.move(str(converted), str(stable_path))
            return stable_path

    def _read_from_doc_antiword(self, path: Path) -> str:
        """
        antiword纯文本提取（表格结构会被拍平）
        """
        import subprocess
        antiword_path = settings.antiword_path
        if not antiword_path or not Path(antiword_path).exists():
            raise ValueError(
                f"无法解析.doc文件：libreoffice_path未配置或转换失败，且antiword_path未配置。"
                f"请至少配置其一。"
            )
        result = subprocess.run(
            [antiword_path, "-m", "UTF-8.txt", str(path)],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            err = result.stderr.decode("utf-8", errors="ignore").strip()
            raise ValueError(f"antiword解析失败: {err}")
        content = result.stdout.decode("utf-8", errors="ignore").strip()
        if not content:
            raise ValueError(f"doc文件{path}无有效文本")
        return content

    def read_from_excel(self, path: Path, suffix: str) -> str:
        """
        从Excel文件读取内容，返回行结构化文本。
        每行格式: 行N → 列名: 值 | 列名: 值
        每行自带表头，便于后续按普通文本切片后每个分片仍可独立检索。
        """
        parts = []
        if suffix == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(filename=str(path), data_only=True, read_only=True)
            try:
                for ws in wb.worksheets:
                    sheet_text = self._sheet_to_row_text(ws.title, ws.iter_rows(values_only=True))
                    if sheet_text:
                        parts.append(sheet_text)
            finally:
                wb.close()
        else:  # .xls
            import xlrd
            book = xlrd.open_workbook(str(path))
            for sheet in book.sheets():
                rows_iter = (
                    [sheet.cell_value(r, c) for c in range(sheet.ncols)]
                    for r in range(sheet.nrows)
                )
                sheet_text = self._sheet_to_row_text(sheet.name, rows_iter)
                if sheet_text:
                    parts.append(sheet_text)

        content = "\n\n".join(parts).strip()
        if not content:
            raise ValueError(f"Excel文件{path}无有效内容")
        return content

    @staticmethod
    def _sheet_to_row_text(sheet_name: str, rows_iter) -> str:
        """
        将一个工作表转为行结构化文本，每行带上表头。
        不同工作表之间由调用方用空行分隔，便于切片器在sheet边界自然切分。
        """
        try:
            headers = next(rows_iter)
        except StopIteration:
            return ""
        headers = [
            str(h).strip() if h is not None and str(h).strip() else f"列{i + 1}"
            for i, h in enumerate(headers)
        ]
        lines = [f"[工作表: {sheet_name}]"]
        row_idx = 2  # 数据从第2行开始（第1行为表头）
        for row in rows_iter:
            if all(c is None or str(c).strip() == "" for c in row):
                row_idx += 1
                continue
            pairs = []
            for h, v in zip(headers, row):
                if v is None or str(v).strip() == "":
                    continue
                pairs.append(f"{h}: {v}")
            if not pairs:
                row_idx += 1
                continue
            lines.append(f"行{row_idx} → " + " | ".join(pairs))
            row_idx += 1
        # 只有标题行没有数据行时返回空
        return "\n".join(lines) if len(lines) > 1 else ""

    def read_from_csv(self, path: Path) -> str:
        """
        从CSV文件读取内容，返回行结构化文本
        """
        import csv
        rows = None
        for encoding in ("utf-8-sig", "utf-8", "gbk"):
            try:
                with open(path, "r", encoding=encoding, newline="") as f:
                    rows = list(csv.reader(f))
                break
            except UnicodeDecodeError:
                continue
        if not rows:
            raise ValueError(f"CSV文件{path}无有效内容")
        return self._sheet_to_row_text(Path(path).stem, iter(rows))


vector_index_service = VectorIndexService()
